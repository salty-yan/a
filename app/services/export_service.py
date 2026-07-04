import os
from datetime import datetime

from app.utils.paths import get_project_root


class ExportService:
    """导出服务，支持导出为 TXT 和 DOCX 格式"""

    def __init__(self):
        pass

    def export_txt(self, question, question_type, spread_name, draw_results, interpretation, save_path=None):
        """导出为文本文件"""
        if not save_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            export_dir = os.path.join(get_project_root(), "exports")
            os.makedirs(export_dir, exist_ok=True)
            save_path = os.path.join(export_dir, f"塔罗占卜报告_{timestamp}.txt")

        try:
            with open(save_path, "w", encoding="utf-8") as f:
                f.write("=" * 50 + "\n")
                f.write("       TarotInsight 塔罗牌占卜报告\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"问    题：{question}\n")
                f.write(f"问题类型：{question_type}\n")
                f.write(f"使用牌阵：{spread_name}\n")
                f.write("\n" + "=" * 50 + "\n")
                f.write("           抽牌结果\n")
                f.write("=" * 50 + "\n\n")

                for i, result in enumerate(draw_results, 1):
                    card = result["card"]
                    is_reversed = result["is_reversed"]
                    direction = "逆位" if is_reversed else "正位"
                    meaning = card["reversed_meaning"] if is_reversed else card["upright_meaning"]

                    f.write(f"--- 第{i}张牌：{result['position']} ---\n")
                    f.write(f"卡牌：{card['name_cn']} ({card['name_en']}) [{direction}]\n")
                    f.write(f"关键词：{'、'.join(card['keywords'])}\n")
                    f.write(f"解读：{meaning}\n\n")

                f.write("=" * 50 + "\n")
                f.write("           综合解读\n")
                f.write("=" * 50 + "\n\n")
                f.write(interpretation)
                f.write("\n\n" + "=" * 50 + "\n")
                f.write("本报告由 TarotInsight 塔罗牌心理探索工具生成\n")
                f.write("仅供心理探索和娱乐参考，不构成专业建议\n")
                f.write("=" * 50 + "\n")

            return save_path, None
        except Exception as e:
            return None, str(e)

    def export_docx(self, question, question_type, spread_name, draw_results, interpretation, save_path=None):
        """导出为 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return None, "python-docx 未安装，请运行: pip install python-docx"

        if not save_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            export_dir = os.path.join(get_project_root(), "exports")
            os.makedirs(export_dir, exist_ok=True)
            save_path = os.path.join(export_dir, f"塔罗占卜报告_{timestamp}.docx")

        try:
            doc = Document()

            # 标题
            title = doc.add_heading("TarotInsight 塔罗牌占卜报告", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 基本信息
            doc.add_heading("基本信息", level=1)
            doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"问    题：{question}")
            doc.add_paragraph(f"问题类型：{question_type}")
            doc.add_paragraph(f"使用牌阵：{spread_name}")

            # 抽牌结果
            doc.add_heading("抽牌结果", level=1)
            for i, result in enumerate(draw_results, 1):
                card = result["card"]
                is_reversed = result["is_reversed"]
                direction = "逆位 ⬇" if is_reversed else "正位 ⬆"
                meaning = card["reversed_meaning"] if is_reversed else card["upright_meaning"]

                doc.add_heading(f"第{i}张牌：{result['position']}", level=2)
                p = doc.add_paragraph()
                p.add_run(f"卡牌：{card['name_cn']} ({card['name_en']}) ").bold = True
                p.add_run(f"[{direction}]")
                doc.add_paragraph(f"关键词：{'、'.join(card['keywords'])}")
                doc.add_paragraph(f"解读：{meaning}")

            # 综合解读
            doc.add_heading("综合解读", level=1)
            doc.add_paragraph(interpretation)

            # 页脚
            doc.add_paragraph("")
            doc.add_paragraph("—" * 30)
            p = doc.add_paragraph("本报告由 TarotInsight 塔罗牌心理探索工具生成")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2 = doc.add_paragraph("仅供心理探索和娱乐参考，不构成专业建议")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(save_path)
            return save_path, None
        except Exception as e:
            return None, str(e)
